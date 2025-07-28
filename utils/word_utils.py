from docx import Document

def load_docx(path):
    return Document(path)

def extract_paragraphs(doc):
    return [p for p in doc.paragraphs if p.text.strip()]

def extract_tables(doc):
    return doc.tables
